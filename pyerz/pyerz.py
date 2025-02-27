# -*- coding: utf-8 -*-
import codecs
import logging
import pkg_resources
import math
from os.path import abspath
try:
    from os import scandir
except ImportError:
    from scandir import scandir

import click
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK


logger = logging.getLogger(__name__)


# 默认源代码文件目录
DEFAULT_INDIRS = ['.']
# 默认支持的代码格式
DEFAULT_EXTS = ['py']
# 默认的注释前缀
DEFAULT_COMMENT_CHARS = (
    '#', '//'
)
DEFAULT_MULTILINE_COMMENT_PAIRS = (
    ('"""', '"""'),  # Python 多行注释
    ("'''", "'''"),  # Python 多行注释
    ('/*', '*/')     # C 风格多行注释
)
ALIGNMENT_MAP = {
    'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
    'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
    'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
}


def del_slash(dirs):
    """
    删除文件夹最后一位的/

    Args:
        dirs: 文件夹列表
    Returns:
        删除之后的文件夹
    """
    no_slash_dirs = []
    for dir_ in dirs:
        if dir_[-1] == '/':
            no_slash_dirs.append(dir_[: -1])
        else:
            no_slash_dirs.append(dir_)
    return no_slash_dirs


class CodeFinder(object):
    """
    给定一个目录，和若干个后缀名，
    递归地遍历该目录，找到该目录下
    所有以这些后缀结束的文件
    """
    def __init__(self, exts=None):
        """
        Args:
            exts: 后缀名，默认为以py结尾
        """
        self.exts = exts if exts else ['py']

    def is_code(self, file):
        for ext in self.exts:
            if file.endswith(ext):
                return True
        return False

    @staticmethod
    def is_hidden_file(file):
        """
        是否是隐藏文件
        """
        return file[0] == '.'

    @staticmethod
    def should_be_excluded(file, excludes=None):
        """
        是否需要略过此文件
        """
        if not excludes:
            return False
        if not isinstance(excludes, list):
            excludes = [excludes]
        should_be_excluded = False
        for exclude in excludes:
            if file.startswith(exclude):
                should_be_excluded = True
                break
        return should_be_excluded

    def find(self, indir, excludes=None):
        """
        给定一个文件夹查找这个文件夹下所有的代码

        Args:
            indir: 需要查到代码的目录
            excludes: 排除文件或目录
        Returns:
            代码文件列表
        """
        files = []
        for entry in scandir(indir):
            # 防止根目录有一些含有非常多文件的隐藏文件夹
            # 例如，.git文件，如果不排除，此程序很难运行
            entry_name = entry.name
            entry_path = abspath(entry.path)
            if self.is_hidden_file(entry_name):
                continue
            if self.should_be_excluded(entry_path, excludes):
                continue
            if entry.is_file():
                if self.is_code(entry_name):
                    files.append(entry_path)
                continue
            for file in self.find(entry_path, excludes=excludes):
                files.append(file)
        logger.debug('在%s目录下找到%d个代码文件.', indir, len(files))
        return files


class CodeWriter(object):
    def __init__(
            self, font_name='宋体',
            font_size=10.5, space_before=0.0,
            space_after=2.3, line_spacing=10.5,
            command_chars=None, multiline_comment_pairs=None,
            document=None, chars_in_line=30, insert_page=False
    ):
        self.font_name = font_name
        self.font_size = font_size
        self.space_before = space_before
        self.space_after = space_after
        self.line_spacing = line_spacing
        self.command_chars = command_chars if command_chars else DEFAULT_COMMENT_CHARS
        self.chars_in_line = chars_in_line * 2
        self.insert_page = insert_page
        self.multiline_comment_pairs = multiline_comment_pairs if multiline_comment_pairs else DEFAULT_MULTILINE_COMMENT_PAIRS
        self.current_comment_pair = None  # 当前正在处理的多行注释对
        self.document = Document(pkg_resources.resource_filename(
            'pyerz', 'template.docx'
        )) if not document else document

        self.total_paragraph_count = 0

    @staticmethod
    def is_blank_line(line):
        """
        判断是否是空行
        """
        return not bool(line)

    def is_comment_line(self, line):
        line = line.lstrip()  # 去除左侧缩进
        
        # 检查是否在多行注释中
        if self.current_comment_pair:
            end_char = self.current_comment_pair[1]
            if end_char in line:
                self.current_comment_pair = None
                return True
            return True
            
        # 检查是否是新的多行注释开始
        for start_char, end_char in self.multiline_comment_pairs:
            if start_char in line:
                if end_char in line[line.index(start_char) + len(start_char):]:
                    # 单行内完成的多行注释
                    return True
                self.current_comment_pair = (start_char, end_char)
                return True
                
        # 检查单行注释
        is_comment = False    # 初始化变量
        for comment_char in self.command_chars:
            if line.startswith(comment_char):
                is_comment = True
                break
        return is_comment

    def write_header(self, title, paragraph_alignment):
        """
        写入页眉
        """
        paragraph = self.document.sections[0].header.paragraphs[0]
        paragraph.alignment = ALIGNMENT_MAP.get(paragraph_alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
        run = paragraph.add_run(title)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        return self

    def write_file(self, file):
        """
        把单个文件添加到程序文档里面
        """
        self.current_comment_pair = None  # 重置多行注释状态
        with codecs.open(file, encoding='utf-8') as fp:
            for line in fp:
                line = line.rstrip()
                if self.is_blank_line(line):
                    continue
                if self.is_comment_line(line):
                    continue
                for i in range(math.ceil(len(line)/self.chars_in_line)):
                    paragraph = self.document.add_paragraph()
                    start = i * self.chars_in_line
                    remain = len(line) - start 
                    if remain >= self.chars_in_line:
                        line_part = line[start:start+self.chars_in_line-1]
                    else:
                        line_part = line[start:]
                    
                    paragraph.paragraph_format.space_before = Pt(self.space_before)
                    paragraph.paragraph_format.space_after = Pt(self.space_after)
                    paragraph.paragraph_format.line_spacing = Pt(self.line_spacing)
                    
                    run = paragraph.add_run(line_part)
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)

                    self.total_paragraph_count += 1

                    # 每 50 行增加一个分页符
                    if self.total_paragraph_count % 50 == 0 and self.insert_page:
                        run.add_break(WD_BREAK.PAGE)
        return self

    def save(self, file):
        self.document.save(file)


@click.command(name='pyerz')
@click.option(
    '-t', '--title', default='软件著作权程序鉴别材料生成器V1.0',
    help='软件名称+版本号，默认为软件著作权程序鉴别材料生成器V1.0，此名称用于生成页眉'
)
@click.option(
    '--entry-file', type=click.Path(exists=True),
    help='入口文件路径，将会被放在文档第一页'
)
@click.option(
    '-i', '--indir', 'indirs',
    multiple=True, type=click.Path(exists=True),
    help='源码所在文件夹，可以指定多个，默认为当前目录'
)
@click.option(
    '-e', '--ext', 'exts',
    multiple=True, help='源代码后缀，可以指定多个，默认为Python源代码'
)
@click.option(
    '-c', '--comment-char', 'comment_chars',
    multiple=True, help='注释字符串，可以指定多个，默认为#、//'
)
@click.option(
    '--multiline-comment-start', 'multiline_starts',
    multiple=True, help='多行注释开始标记，需要与结束标记一一对应'
)
@click.option(
    '--multiline-comment-end', 'multiline_ends',
    multiple=True, help='多行注释结束标记，需要与开始标记一一对应'
)
@click.option(
    '--font-name', default='宋体',
    help='字体，默认为宋体'
)
@click.option(
    '--font-size', default=10.5,
    type=click.FloatRange(min=1.0),
    help='字号，默认为五号，即10.5号'
)
@click.option(
    '--space-before', default=0.0,
    type=click.FloatRange(min=0.0),
    help='段前间距，默认为0'
)
@click.option(
    '--paragraph-alignment', default='center',
    type=click.Choice(['left', 'center', 'right']),
    help='段对齐方式，默认为居中对齐'
)
@click.option(
    '--space-after', default=2.3,
    type=click.FloatRange(min=0.0),
    help='段后间距，默认为2.3'
)
@click.option(
    '--line-spacing', default=10.5,
    type=click.FloatRange(min=0.0),
    help='行距，默认为固定值10.5'
)
@click.option(
    '--chars-in-line', default=30, 
    type=click.IntRange(min=1), 
    help='一行的字符数，中文字符（2字节），默认为30'
)
@click.option(
    '--exclude', 'excludes',
    multiple=True, type=click.Path(exists=True),
    help='需要排除的文件或路径，可以指定多个'
)
@click.option(
    '-o', '--outfile', default='code.docx',
    type=click.Path(exists=False),
    help='输出文件（docx格式），默认为当前目录的code.docx'
)
@click.option('-p', '--insert-page', is_flag=True, help='每50行插入一个分页符')
@click.option('-v', '--verbose', is_flag=True, help='打印调试信息')
def main(
        title, indirs, exts, entry_file,
        comment_chars, multiline_starts, multiline_ends,
        font_name, font_size, space_before,
        space_after, line_spacing,
        chars_in_line, paragraph_alignment,
        excludes, outfile, insert_page, verbose,
):
    if not indirs:
        indirs = DEFAULT_INDIRS
    if not exts:
        exts = DEFAULT_EXTS
    if not comment_chars:
        comment_chars = DEFAULT_COMMENT_CHARS
    if verbose:
        logging.basicConfig(level=logging.DEBUG)

    # 处理多行注释配置
    multiline_pairs = None
    if multiline_starts and multiline_ends and len(multiline_starts) == len(multiline_ends):
        multiline_pairs = list(zip(multiline_starts, multiline_ends))

    # 第零步，把所有的路径都转换为绝对路径
    indirs = [abspath(indir) for indir in indirs]
    excludes = del_slash(
        [abspath(exclude) for exclude in excludes] if excludes else []
    )

    # 第一步，查找代码文件
    finder = CodeFinder(exts)
    files = [file for indir in indirs for file in finder.find(
        indir, excludes=excludes
    )]

    if entry_file:
        entry_path = abspath(entry_file)
        # 如果入口文件在列表中，先移除它
        if entry_path in files:
            files.remove(entry_path)
        # 将入口文件放在列表首位
        files.insert(0, entry_path)

    # 第二步，逐个把代码文件写入到docx中
    writer = CodeWriter(
        command_chars=comment_chars,
        multiline_comment_pairs=multiline_pairs,
        font_name=font_name,
        font_size=font_size,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
        chars_in_line=chars_in_line, 
        insert_page=insert_page
    )
    writer.write_header(title, paragraph_alignment)
    for file in files:
        writer.write_file(file)
    writer.save(outfile)
    return 0


if __name__ == '__main__':
    main()
